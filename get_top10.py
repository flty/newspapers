# -*- coding:utf-8 -*- 

import datetime
import urllib2
import re
import os
import lxml.html
from lxml import etree
from docx import Document
from docx.shared import Inches


def today():
    day = datetime.datetime.today().date()
    return str(day) 

def get_top10(date=None, retry_count=3, pause=0.001):
    """
        
    """
    
    url = 'http://paperpost.people.com.cn/rmrb-%s.html' %(date)
    request = urllib2.Request(url)
    text = urllib2.urlopen(request,timeout=10).read()
    text = text.decode('utf-8')
    top10_url = re.compile(r'<a href=\"(.*?)\"')
    top10_url = top10_url.findall(text)

    for next_url in range(0,10):
    	html = lxml.html.parse(top10_url[next_url])
        res = html.xpath('//div[@id=\"ozoom\"]')[0]
        res = res.xpath('//p')
        H1 = html.xpath('//h1')[0].text
        newdoc = Document() #添加新文件
        newdoc.add_heading(H1, 0)
        for node in res:        	
        	sarr = node.text
        	if sarr :
        		newdoc.add_paragraph(sarr)
        dir_to_save = r'G:/rmrb_top10/' + date + r'/'
        newdoc.save(dir_to_save + (r'%s.docx') %H1)#另存为

if __name__ == '__main__' :
	print today()
	dir_to_save = r'G:/rmrb_top10/' + today() + r'/'
	if not os.path.isdir(dir_to_save):
		os.mkdir(dir_to_save)
	get_top10(today())